from docx import Document
from lxml import etree
import qrcode
import random




def generate_eq():
    # Generer likning på formen: a + bx/f = c + dx
    low = -10
    high = 10
    # enkle
    a = random.choice([i for i in range(low,high) if i not in [0]])
    b = random.choice([i for i in range(low,high) if i not in [0]])
    x = random.choice([i for i in range(low,high) if i not in [0]])

    f = random.choice([i for i in range(low,high) if i not in [-1,0,1,b,abs(b*-1)]])

    c = random.choice([i for i in range(low,high) if i not in [0]])
    d = random.choice([i for i in range(low,high) if i not in [0,b]])
    
    vs = []
    for e in ["a","b*x/f"]:
        e_str = e.replace("a",str(a)).replace("b",str(b)).replace("f",str(f)).replace("x",str(x))
        vs.append(e_str)
        
    
    vs_val_string = ''
    for e in vs:
        vs_val_string += "(" +e +")+"
    
    vs_val_string = vs_val_string[:-1] 
    vs_val_string = vs_val_string.replace("(=)","=")
    vs_val_string = vs_val_string.replace("+=","=")
    vs_val_string = vs_val_string.replace("=+","=")
    vs_val_string = vs_val_string.replace("+-","-")
    vs_val_string = vs_val_string.replace("-1*","")
    vs_val_string = vs_val_string.replace("1*","")
    
   # print(vs_val_string)
    vs_value = eval(vs_val_string)
    #print(vs_value)
    
    dx_val = eval("d*x")
    c_val = vs_value - dx_val
    if c_val % 1 != 0 or c_val == 0:
        return(False,False)
        
    #print(a,b,c,f,x)
    if random.randint(0,1):
        vs_string = f'{a};{b}x/{f}'
    else:
        vs_string = f'{b}x/{f};{a}'
        
    if random.randint(0,1):    
        hs_string = f'{c_val};{d}x'
    else:
        hs_string = f'{d}x;{c_val}'
        
    if random.randint(0,1):
        hs_string,vs_string = vs_string,hs_string
    
    eq_string = vs_string + ";=;" + hs_string
    eq_string = eq_string.replace("+-","-")
    eq_string = eq_string.replace("1x","x")
    eq_string = eq_string.replace(".0","")
    eq_string = eq_string.replace(".",",")
    
    
    print(eq_string, x)
    
    return(eq_string,x)


#https://elsenaju.eu/mathml/MathML-Examples.htm
def math_to_word(eq):
    tree = etree.fromstring(eq)
    xslt = etree.parse("MML2OMML.XSL")
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    return new_dom.getroot()

def eq_to_MathML(eqList):
    MathML_string = ''
    first = True         # Brukes for å forhindre at det blir lagt til + før første ledd
    negative = False     # Brukes for å forhindre at det blir lagt til + før negative ledd
    eqList = eqList.split(";")
    
    for ledd in eqList:
        if ledd.startswith("-") and "/" not in ledd:
                  negative = True
                  
        if ledd != "=" and not first and not negative:
            MathML_string += '<mo>+</mo>'
            
        first = False
        negative = False
      
        if "/" in ledd:
            teller,nevner = ledd.split("/")
            MathML_string += f'<mfrac><mrow><mi>{teller}</mi></mrow><mi>{nevner}</mi></mfrac>'
        else:  
            MathML_string += f'<mi>{ledd}</mi>'
            
        if ledd == "=":
            first = True
            
        # Legg til støtte for potenser: https://developer.mozilla.org/en-US/docs/Web/MathML/Element/msup
        
    return MathML_string


mathstr_start = '<math xmlns="http://www.w3.org/1998/Math/MathML">'
mathstr_end = '</math>'

doc = Document()
table = doc.add_table(rows=8, cols=3, style="Table Grid")

nr = 0
for row in table.rows:
    for cell in row.cells:
        validEQ = False
        while not validEQ:
            eq_str,fasit = generate_eq()
            if eq_str:
                validEQ = True
        p = cell.add_paragraph()
        p._element.append(math_to_word(mathstr_start + eq_to_MathML(eq_str) + mathstr_end))
        
        
        img = qrcode.make("x = "+ str(fasit))
        img.save(f'qr{nr}.png')
        
        p = cell.add_paragraph()
        p.alignment = 1
        run = p.add_run()
        run.add_picture(f'qr{nr}.png', width = 500000, height = 500000)
        nr += 1

doc.save('test2.docx')

